import os
import streamlit as st
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage
from docx import Document
from io import BytesIO
import base64



st.set_page_config(
    page_title="Gerador de Contratos",
    page_icon="./assets/icons/icon.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Função para codificar a imagem em base64
def get_base64_image(image_path):
    with open(image_path, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

# Configuração inicial da página
def apply_custom_styles():
    image_path = './assets/images/logo.png'  # Caminho para o seu logo
    encoded_image = get_base64_image(image_path)
    st.markdown(f"""
    <style>
        [data-testid="stSidebar"]::before {{
            content: "";
            display: block;
            height: 130px;
            background-image: url('data:image/png;base64,{encoded_image}');
            background-size: contain;
            background-repeat: no-repeat;
            background-position: center;
            margin-top:40px;
        }}
        [data-testid="stSidebar"] .block-container {{
            padding-top: 20px;
        }}
        [data-testid="stSidebar"] {{
            background-color: #1C1C1C;
            color: white;
        }}
        .stButton>button {{
            background-color: #587472;
            color: white;
            border-radius: 8px;
            font-size: 16px;
            margin-bottom: 10px;
        }}
        .stTabs [role=tab] {{
            background-color: #40444B;
            color: white;
            padding: 10px;
            font-size: 18px;
        }}
        .stTabs [role=tab][aria-selected="true"] {{
            background-color: #587472;
        }}
    </style>
    """, unsafe_allow_html=True)

apply_custom_styles()

# Mapeamento de agentes para subdiretórios
agentes_map = {
    "Prestação de Serviços": "prestacao_servicos",
    "Parcerias & Coproduções": "parcerias_coproducoes",
    "Contratação de Colaboradores": "contratacao_colaboradores",
    "Confidencialidade e Não Concorrência": "confidencialidade",
    "Políticas de Privacidade": "termos_politicas",
    "Propriedade Intelectual": "propriedade_intelectual",
    "Mentorias & Masterminds": "mentorias_masterminds",
    "Acordos Societários": "acordos_societarios",
}

BASE_EXEMPLOS_DIR = "exemplos"

# Sidebar para seleção de agentes
st.sidebar.title("Selecione o tipo de contrato")
agente_selecionado = st.sidebar.selectbox("Agentes", list(agentes_map.keys()))
st.header('Gerador de Contratos', divider='green')

# Função para extrair texto de arquivos DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

# Função para carregar exemplos
def identificar_tipos_contratos(pasta_exemplo):
    subtipos = {}
    if os.path.exists(pasta_exemplo):
        arquivos_exemplo = [
            f for f in os.listdir(pasta_exemplo) if f.endswith(".docx")
        ]
        for arquivo in arquivos_exemplo:
            caminho_arquivo = os.path.join(pasta_exemplo, arquivo)
            texto = extract_text_from_docx(caminho_arquivo)
            subtipos[arquivo] = texto
    return subtipos

# Obter a pasta correspondente
pasta_agente = agentes_map.get(agente_selecionado)
pasta_exemplo = os.path.join(BASE_EXEMPLOS_DIR, pasta_agente)
subtipos_contratos = identificar_tipos_contratos(pasta_exemplo)

st.header("1. Escolha o tipo de contrato")

if subtipos_contratos:
    tipo_contrato_selecionado = st.selectbox(
        "Selecione o tipo de contrato desejado:", list(subtipos_contratos.keys())
    )
else:
    st.warning("Nenhum exemplo de contrato encontrado para este agente.")
    st.stop()

st.header("2. Responda às perguntas para personalizar o contrato")

texto_exemplo = subtipos_contratos[tipo_contrato_selecionado]

# Configurar o modelo GPT-4 via LangChain
llm = ChatOpenAI(
    model="gpt-4o-mini",
    openai_api_key=st.secrets["OPENAI_API_KEY"],  # Chave de API protegida em secrets
)

@st.cache_data
def gerar_perguntas(prompt):
    response = llm(messages=[HumanMessage(content=prompt)])
    return response.content.split("\n")

prompt_perguntas = f"""
Você é um especialista jurídico com a tarefa de personalizar contratos. Analise o seguinte contrato de exemplo:

{texto_exemplo}

Identifique os 15 campos variáveis mais importantes e crie uma lista de perguntas que devem ser feitas ao usuário para preencher todas as informações necessárias para personalizar completamente o contrato, sem campos faltantes.

As perguntas devem:
1. Ser claras e objetivas.
2. Abranger informações críticas como nomes, valores, prazos, responsabilidades, condições específicas, e qualquer outro detalhe indispensável.
3. Focar apenas nos elementos essenciais para evitar redundâncias.
4. Estar formatadas como uma lista numerada.

Por favor, responda com apenas as 15 perguntas.
"""

perguntas = gerar_perguntas(prompt_perguntas)

respostas = {}

with st.form("formulario_perguntas"):
    st.write("Por favor, preencha as informações abaixo:")
    for pergunta in perguntas:
        if pergunta.strip():
            respostas[pergunta] = st.text_input(pergunta)
    submit = st.form_submit_button("Enviar Respostas")

st.header("3. Contrato Gerado")

if submit:
    st.success("Respostas salvas com sucesso!")
    prompt_contrato = f"""
Você é um assistente jurídico e recebeu um contrato de exemplo e as respostas fornecidas pelo usuário para campos variáveis. Sua tarefa é gerar um contrato completo e personalizado.

Contrato de exemplo para ser seguido:
{texto_exemplo}

Respostas do usuário:
{', '.join([f"{k}: {v}" for k, v in respostas.items()])}

Requisitos:
1. Use a estrutura e o texto do contrato de exemplo como base.
2. Substitua os campos variáveis pelas respostas do usuário.
3. Certifique-se de que o contrato resultante esteja completo, claro e juridicamente correto, sem campos faltantes.
4. Mantenha o tom profissional e a formatação consistente.

Gere o contrato completo e personalizado abaixo:
"""
    response_contrato = llm(messages=[HumanMessage(content=prompt_contrato)])
    contrato_gerado = response_contrato.content
    st.text_area("Contrato Gerado", contrato_gerado, height=400)
      # Somente a opção de baixar como DOCX
    doc = Document()
    doc.add_paragraph(contrato_gerado)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    st.download_button(
        "Baixar como DOCX",
        data=doc_io,
        file_name="contrato_gerado.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )